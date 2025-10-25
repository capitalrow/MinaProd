"""
Transcript API Routes
REST API endpoints for transcript export, editing, and speaker management.
Part of Phase 2: Transcript Experience Enhancement (T2.3, T2.5, T2.6)
"""

from flask import Blueprint, request, jsonify, send_file
from flask_login import login_required, current_user
from models import db, Meeting, Session, Segment, SegmentComment
from datetime import datetime
from io import BytesIO
import json
import logging

logger = logging.getLogger(__name__)

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


api_transcript_bp = Blueprint('api_transcript', __name__, url_prefix='/api/meetings')


# ============================================
# Export Endpoints (T2.3)
# ============================================

@api_transcript_bp.route('/<int:meeting_id>/export/<format>', methods=['GET'])
@login_required
def export_transcript(meeting_id, format):
    """Export transcript in various formats: txt, docx, pdf, json."""
    # Get meeting and verify ownership
    meeting = db.session.query(Meeting).filter_by(
        id=meeting_id,
        workspace_id=current_user.workspace_id
    ).first()
    
    if not meeting:
        return jsonify({'success': False, 'message': 'Meeting not found'}), 404
    
    # Get transcript segments
    segments = []
    if meeting.session_id:
        session = db.session.query(Session).filter_by(external_id=meeting.session_id).first()
        if session:
            segments = db.session.query(Segment).filter_by(
                session_id=session.id,
                kind='final'
            ).order_by(Segment.start_ms.asc()).all()
    
    if not segments:
        return jsonify({'success': False, 'message': 'No transcript available'}), 404
    
    # Generate export based on format
    if format == 'txt':
        return export_txt(meeting, segments)
    elif format == 'docx' and DOCX_AVAILABLE:
        return export_docx(meeting, segments)
    elif format == 'pdf' and PDF_AVAILABLE:
        return export_pdf(meeting, segments)
    elif format == 'json':
        return export_json(meeting, segments)
    else:
        return jsonify({'success': False, 'message': f'Export format "{format}" not supported'}), 400


def export_txt(meeting, segments):
    """Export transcript as plain text file."""
    # Build text content
    lines = []
    lines.append(f"Meeting: {meeting.title}")
    lines.append(f"Date: {meeting.created_at.strftime('%Y-%m-%d %H:%M')}")
    lines.append(f"Duration: {meeting.duration_minutes or 'Unknown'} minutes")
    lines.append("=" * 60)
    lines.append("")
    
    current_speaker = None
    for segment in segments:
        speaker = getattr(segment, 'speaker_name', None) or getattr(segment, 'speaker_id', 'Speaker')
        timestamp = segment.start_time_formatted if hasattr(segment, 'start_time_formatted') else ""
        
        # Add speaker header if changed
        if speaker != current_speaker:
            if current_speaker is not None:
                lines.append("")
            lines.append(f"[{timestamp}] {speaker}:")
            current_speaker = speaker
        
        lines.append(f"  {segment.text}")
    
    # Create file
    content = "\n".join(lines)
    buffer = BytesIO(content.encode('utf-8'))
    buffer.seek(0)
    
    filename = f"{meeting.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt"
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='text/plain'
    )


def export_docx(meeting, segments):
    """Export transcript as Microsoft Word document."""
    if not DOCX_AVAILABLE:
        return jsonify({'success': False, 'message': 'DOCX export not available'}), 503
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading(meeting.title, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add metadata
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_para.add_run(f"Date: {meeting.created_at.strftime('%Y-%m-%d %H:%M')}\n")
    meta_para.add_run(f"Duration: {meeting.duration_minutes or 'Unknown'} minutes")
    
    doc.add_paragraph()  # Spacer
    
    # Add transcript segments
    current_speaker = None
    for segment in segments:
        speaker = getattr(segment, 'speaker_name', None) or getattr(segment, 'speaker_id', 'Speaker')
        timestamp = segment.start_time_formatted if hasattr(segment, 'start_time_formatted') else ""
        
        # Add speaker header if changed
        if speaker != current_speaker:
            if current_speaker is not None:
                doc.add_paragraph()  # Spacer between speakers
            
            speaker_para = doc.add_paragraph()
            speaker_run = speaker_para.add_run(f"[{timestamp}] {speaker}:")
            speaker_run.bold = True
            speaker_run.font.size = Pt(12)
            speaker_run.font.color.rgb = RGBColor(99, 102, 241)  # Primary color
            current_speaker = speaker
        
        # Add segment text
        text_para = doc.add_paragraph(segment.text)
        text_para.paragraph_format.left_indent = Pt(20)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{meeting.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


def export_pdf(meeting, segments):
    """Export transcript as PDF document."""
    if not PDF_AVAILABLE:
        return jsonify({'success': False, 'message': 'PDF export not available'}), 503
    
    # Create PDF buffer
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=inch, leftMargin=inch,
                           topMargin=inch, bottomMargin=inch)
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#6366f1',
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    meta_style = ParagraphStyle(
        'MetaStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor='#6b7280',
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    speaker_style = ParagraphStyle(
        'SpeakerStyle',
        parent=styles['Normal'],
        fontSize=12,
        textColor='#6366f1',
        bold=True,
        spaceAfter=6
    )
    
    text_style = ParagraphStyle(
        'TextStyle',
        parent=styles['Normal'],
        fontSize=11,
        leftIndent=20,
        spaceAfter=10
    )
    
    # Build story
    story = []
    
    # Title
    story.append(Paragraph(meeting.title, title_style))
    
    # Metadata
    meta_text = f"Date: {meeting.created_at.strftime('%Y-%m-%d %H:%M')} | Duration: {meeting.duration_minutes or 'Unknown'} minutes"
    story.append(Paragraph(meta_text, meta_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Transcript segments
    current_speaker = None
    for segment in segments:
        speaker = getattr(segment, 'speaker_name', None) or getattr(segment, 'speaker_id', 'Speaker')
        timestamp = segment.start_time_formatted if hasattr(segment, 'start_time_formatted') else ""
        
        # Add speaker header if changed
        if speaker != current_speaker:
            if current_speaker is not None:
                story.append(Spacer(1, 0.15*inch))
            
            speaker_text = f"[{timestamp}] {speaker}:"
            story.append(Paragraph(speaker_text, speaker_style))
            current_speaker = speaker
        
        # Add segment text
        story.append(Paragraph(segment.text, text_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    filename = f"{meeting.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/pdf'
    )


def export_json(meeting, segments):
    """Export transcript as structured JSON."""
    # Build JSON structure
    data = {
        'meeting': {
            'id': meeting.id,
            'title': meeting.title,
            'date': meeting.created_at.isoformat() if meeting.created_at else None,
            'duration_minutes': meeting.duration_minutes,
            'status': meeting.status,
        },
        'transcript': {
            'segments': []
        },
        'metadata': {
            'exported_at': datetime.utcnow().isoformat(),
            'total_segments': len(segments),
            'format_version': '1.0'
        }
    }
    
    for segment in segments:
        data['transcript']['segments'].append({
            'id': segment.id,
            'speaker': getattr(segment, 'speaker_name', None) or getattr(segment, 'speaker_id', 'Speaker'),
            'text': segment.text,
            'start_ms': segment.start_ms,
            'end_ms': segment.end_ms,
            'start_time_formatted': segment.start_time_formatted if hasattr(segment, 'start_time_formatted') else None,
            'confidence': segment.avg_confidence,
            'is_final': segment.is_final
        })
    
    # Create JSON response
    buffer = BytesIO(json.dumps(data, indent=2).encode('utf-8'))
    buffer.seek(0)
    
    filename = f"{meeting.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.json"
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/json'
    )


# ============================================
# Edit Endpoints (T2.5)
# ============================================

@api_transcript_bp.route('/<int:meeting_id>/segments/<int:segment_id>', methods=['PATCH'])
@login_required
def update_segment(meeting_id, segment_id):
    """Update transcript segment text (inline editing)."""
    # Verify meeting ownership
    meeting = db.session.query(Meeting).filter_by(
        id=meeting_id,
        workspace_id=current_user.workspace_id
    ).first()
    
    if not meeting:
        return jsonify({'success': False, 'message': 'Meeting not found'}), 404
    
    # Get segment
    segment = db.session.query(Segment).filter_by(id=segment_id).first()
    
    if not segment:
        return jsonify({'success': False, 'message': 'Segment not found'}), 404
    
    # Get new text and version from request
    data = request.get_json()
    new_text = data.get('text', '').strip()
    client_version = data.get('version')  # Client sends current version for conflict detection
    
    if not new_text:
        return jsonify({'success': False, 'message': 'Text is required'}), 400
    
    # CROWN+ Version Control: Check for concurrent edit conflicts
    if client_version is not None and hasattr(segment, 'version'):
        if segment.version != client_version:
            return jsonify({
                'success': False,
                'conflict': True,
                'message': f'This segment was edited by another user. Please review their changes.',
                'current_version': segment.version,
                'current_text': segment.text,
                'your_version': client_version,
                'your_text': new_text
            }), 409  # 409 Conflict
    
    # Update segment
    old_text = segment.text
    segment.text = new_text
    
    # CROWN+ Version Control: Increment version on edit
    if hasattr(segment, 'version'):
        segment.version += 1
    
    # Add edit metadata
    if hasattr(segment, 'edited_at'):
        segment.edited_at = datetime.utcnow()
    if hasattr(segment, 'edited_by_id'):
        segment.edited_by_id = current_user.id
    
    try:
        # Get session for event tracking BEFORE commit
        session = db.session.query(Session).filter_by(id=segment.session_id).first()
        
        # Commit database changes
        db.session.commit()
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    
    # CROWN+ Event Chain: Log edit_commit event AFTER successful commit
    # Separate try/except to ensure API returns success even if event logging fails
    event_logged = False
    if session:
        try:
            from services.event_tracking import get_event_tracker
            event_tracker = get_event_tracker()
            event_tracker.edit_commit(
                session=session,
                segment_id=segment.id,
                old_text=old_text,
                new_text=new_text,
                user_id=current_user.id
            )
            event_logged = True
        except Exception as event_error:
            # Edit succeeded but event logging failed - log warning but don't fail the request
            logger.warning(f"⚠️ CROWN+ WARNING: edit_commit event logging failed for segment {segment.id}: {event_error}")
            # Continue - edit was successful even if event logging failed
    
    return jsonify({
        'success': True,
        'message': 'Segment updated successfully',
        'segment': {
            'id': segment.id,
            'text': segment.text,
            'old_text': old_text,
            'version': segment.version if hasattr(segment, 'version') else None,
            'edited_at': segment.edited_at.isoformat() if hasattr(segment, 'edited_at') and segment.edited_at else None
        },
        'event_logged': event_logged  # Indicate if event was logged for debugging
    })


# ============================================
# Speaker Management Endpoints (T2.6)
# ============================================

@api_transcript_bp.route('/<int:meeting_id>/speakers', methods=['PATCH'])
@login_required
def update_speaker_name(meeting_id):
    """Update speaker name across all segments."""
    # Verify meeting ownership
    meeting = db.session.query(Meeting).filter_by(
        id=meeting_id,
        workspace_id=current_user.workspace_id
    ).first()
    
    if not meeting:
        return jsonify({'success': False, 'message': 'Meeting not found'}), 404
    
    # Get speaker names from request
    data = request.get_json()
    old_name = data.get('oldName', '').strip()
    new_name = data.get('newName', '').strip()
    
    if not old_name or not new_name:
        return jsonify({'success': False, 'message': 'Both oldName and newName are required'}), 400
    
    # Get session
    if not meeting.session_id:
        return jsonify({'success': False, 'message': 'No transcript available'}), 404
    
    session = db.session.query(Session).filter_by(external_id=meeting.session_id).first()
    if not session:
        return jsonify({'success': False, 'message': 'Session not found'}), 404
    
    # Update all segments with this speaker
    # Note: This assumes speaker info is stored in segment metadata
    # If speaker identification is not yet in the Segment model, this will need database migration
    updated_count = 0
    
    segments = db.session.query(Segment).filter_by(session_id=session.id).all()
    for segment in segments:
        # Check if segment has speaker_name or speaker_id attribute
        if hasattr(segment, 'speaker_name') and segment.speaker_name == old_name:
            segment.speaker_name = new_name
            updated_count += 1
        elif hasattr(segment, 'speaker_id') and segment.speaker_id == old_name:
            segment.speaker_id = new_name
            updated_count += 1
    
    try:
        if updated_count > 0:
            db.session.commit()
            return jsonify({
                'success': True,
                'message': f'Updated {updated_count} segments',
                'updated_count': updated_count
            })
        else:
            return jsonify({
                'success': True,
                'message': 'No segments found with this speaker name',
                'updated_count': 0
            })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


@api_transcript_bp.route('/<int:meeting_id>/speakers', methods=['GET'])
@login_required
def get_speakers(meeting_id):
    """Get list of speakers in the meeting."""
    # Verify meeting ownership
    meeting = db.session.query(Meeting).filter_by(
        id=meeting_id,
        workspace_id=current_user.workspace_id
    ).first()
    
    if not meeting:
        return jsonify({'success': False, 'message': 'Meeting not found'}), 404
    
    # Get session
    if not meeting.session_id:
        return jsonify({'success': False, 'message': 'No transcript available'}), 404
    
    session = db.session.query(Session).filter_by(external_id=meeting.session_id).first()
    if not session:
        return jsonify({'success': False, 'message': 'Session not found'}), 404
    
    # Get unique speakers
    segments = db.session.query(Segment).filter_by(session_id=session.id).all()
    
    speakers = {}
    for segment in segments:
        speaker_name = getattr(segment, 'speaker_name', None) or getattr(segment, 'speaker_id', 'Speaker')
        
        if speaker_name not in speakers:
            speakers[speaker_name] = {
                'name': speaker_name,
                'segment_count': 0,
                'total_words': 0
            }
        
        speakers[speaker_name]['segment_count'] += 1
        speakers[speaker_name]['total_words'] += len(segment.text.split())
    
    return jsonify({
        'success': True,
        'speakers': list(speakers.values())
    })


# ============================================
# Highlight Endpoints (T2.7)
# ============================================

@api_transcript_bp.route('/<int:meeting_id>/segments/<int:segment_id>/highlight', methods=['PATCH'])
@login_required
def update_segment_highlight(meeting_id, segment_id):
    """Update segment highlight color."""
    # Verify meeting ownership
    meeting = db.session.query(Meeting).filter_by(
        id=meeting_id,
        workspace_id=current_user.workspace_id
    ).first()
    
    if not meeting:
        return jsonify({'success': False, 'message': 'Meeting not found'}), 404
    
    # Get segment
    segment = db.session.query(Segment).filter_by(id=segment_id).first()
    
    if not segment:
        return jsonify({'success': False, 'message': 'Segment not found'}), 404
    
    # Get highlight color from request
    data = request.get_json()
    highlight_color = data.get('highlightColor', None)
    
    # Store highlight in segment metadata or separate table
    # For now, we'll use a simple approach with metadata
    # In production, you'd want a separate highlights table
    if hasattr(segment, 'metadata'):
        if segment.metadata is None:
            segment.metadata = {}
        segment.metadata['highlight_color'] = highlight_color
    
    try:
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Highlight updated successfully',
            'highlight_color': highlight_color
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


# ============================================
# SegmentComment Endpoints (T2.8)
# ============================================

@api_transcript_bp.route('/<int:meeting_id>/segments/<int:segment_id>/comments', methods=['GET'])
@login_required
def get_segment_comments(meeting_id, segment_id):
    """Get comments for a specific segment."""
    # Verify meeting ownership
    meeting = db.session.query(Meeting).filter_by(
        id=meeting_id,
        workspace_id=current_user.workspace_id
    ).first()
    
    if not meeting:
        return jsonify({'success': False, 'message': 'Meeting not found'}), 404
    
    # Get segment
    segment = db.session.query(Segment).filter_by(id=segment_id).first()
    
    if not segment:
        return jsonify({'success': False, 'message': 'Segment not found'}), 404
    
    # Query comments from database (only top-level comments, includes replies)
    comments = db.session.query(SegmentComment).filter_by(
        segment_id=segment_id,
        parent_id=None  # Only get top-level comments
    ).order_by(SegmentComment.created_at.asc()).all()
    
    comments_data = [comment.to_dict(include_replies=True) for comment in comments]
    
    return jsonify({
        'success': True,
        'comments': comments_data
    })


@api_transcript_bp.route('/<int:meeting_id>/segments/<int:segment_id>/comments', methods=['POST'])
@login_required
def add_segment_comment(meeting_id, segment_id):
    """Add a comment to a segment."""
    # Verify meeting ownership
    meeting = db.session.query(Meeting).filter_by(
        id=meeting_id,
        workspace_id=current_user.workspace_id
    ).first()
    
    if not meeting:
        return jsonify({'success': False, 'message': 'Meeting not found'}), 404
    
    # Get segment
    segment = db.session.query(Segment).filter_by(id=segment_id).first()
    
    if not segment:
        return jsonify({'success': False, 'message': 'Segment not found'}), 404
    
    # Get comment text from request
    data = request.get_json()
    text = data.get('text', '').strip()
    parent_id = data.get('parent_id')  # For threaded replies
    
    if not text:
        return jsonify({'success': False, 'message': 'SegmentComment text is required'}), 400
    
    # Validate parent_id if provided (for reply threading)
    if parent_id:
        parent_comment = db.session.query(SegmentComment).filter_by(id=parent_id).first()
        if not parent_comment or parent_comment.segment_id != segment_id:
            return jsonify({'success': False, 'message': 'Invalid parent comment'}), 400
    
    # Create and save comment to database
    try:
        comment = SegmentComment(
            segment_id=segment_id,
            user_id=current_user.id,
            text=text,
            parent_id=parent_id
        )
        db.session.add(comment)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'SegmentComment added successfully',
            'comment': comment.to_dict(include_replies=False)
        }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

# --- Interactive Transcript Endpoints ---
@api_transcript_bp.route("/api/segment/<int:segment_id>/highlight", methods=["POST"])
@login_required
def highlight_segment(segment_id):
    data = request.get_json() or {}
    seg = Segment.query.get_or_404(segment_id)
    seg.is_highlighted = bool(data.get("highlighted", False))
    seg.highlight_color = data.get("color", "yellow")
    db.session.commit()
    return jsonify({"status": "ok", "segment_id": segment_id, "highlighted": seg.is_highlighted})

@api_transcript_bp.route("/api/segment/<int:segment_id>/comment", methods=["POST"])
@login_required
def add_comment(segment_id):
    data = request.get_json() or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"status": "error", "message": "Empty comment"}), 400
    new_comment = SegmentComment(segment_id=segment_id, user_id=current_user.id, text=text, created_at=datetime.utcnow())
    db.session.add(new_comment)
    db.session.commit()
    return jsonify({"status": "ok", "comment": {"id": new_comment.id, "text": new_comment.text}})