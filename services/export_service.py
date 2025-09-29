"""
Advanced Export Service for Mina.

This module provides comprehensive export functionality for meeting transcripts,
summaries, and analytics in multiple formats including PDF, DOCX, and Markdown.
Supports both legacy single-session exports and advanced multi-session exports.
"""

import io
import json
import logging
from typing import Optional, Dict, List, Any, Union
from datetime import datetime
from dataclasses import dataclass
from enum import Enum
from services.session_service import SessionService

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, blue
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class ExportService:
    """Service class for exporting session data to various formats."""
    
    @staticmethod
    def session_to_markdown(session_id: int) -> Optional[str]:
        """
        Convert a session to Markdown format.
        
        Args:
            session_id: Database session ID
            
        Returns:
            Markdown string or None if session not found
        """
        # Get session detail with segments
        session_detail = SessionService.get_session_detail(session_id)
        if not session_detail:
            return None
        
        session = session_detail['session']
        segments = session_detail['segments']
        
        # Start building markdown
        md_lines = []
        
        # Header
        md_lines.append(f"# {session['title']}")
        md_lines.append("")
        
        # Session metadata
        md_lines.append("## Session Information")
        md_lines.append("")
        md_lines.append(f"- **Session ID:** `{session['external_id']}`")
        md_lines.append(f"- **Status:** {session['status'].title()}")
        
        if session.get('started_at'):
            started = session['started_at'].replace('T', ' ')[:16]
            md_lines.append(f"- **Started:** {started}")
        
        if session.get('completed_at'):
            completed = session['completed_at'].replace('T', ' ')[:16]
            md_lines.append(f"- **Completed:** {completed}")
        
        if session.get('locale'):
            md_lines.append(f"- **Language:** {session['locale']}")
        
        md_lines.append(f"- **Total Segments:** {session['segments_count']}")
        
        # Device info if available
        if session.get('device_info'):
            md_lines.append("- **Device Information:**")
            for key, value in session['device_info'].items():
                md_lines.append(f"  - {key.title()}: {value}")
        
        md_lines.append("")
        
        # Transcript section
        if segments:
            md_lines.append("## Transcript")
            md_lines.append("")
            
            # Separate final and interim segments
            final_segments = [s for s in segments if s.get('is_final')]
            interim_segments = [s for s in segments if not s.get('is_final')]
            
            # Export final segments first (main transcript)
            if final_segments:
                md_lines.append("### Final Transcript")
                md_lines.append("")
                
                for segment in final_segments:
                    timestamp = segment.get('start_time_formatted', '00:00')
                    confidence = segment.get('avg_confidence')
                    text = segment.get('text', '').strip()
                    
                    if confidence:
                        conf_display = f" *(confidence: {confidence*100:.1f}%)*"
                    else:
                        conf_display = ""
                    
                    md_lines.append(f"**[{timestamp}]** {text}{conf_display}")
                    md_lines.append("")
            
            # Add interim segments if no final segments exist
            if not final_segments and interim_segments:
                md_lines.append("### Interim Results")
                md_lines.append("")
                md_lines.append("*Note: This session contains only interim transcription results.*")
                md_lines.append("")
                
                for segment in interim_segments[:10]:  # Limit interim to first 10
                    timestamp = segment.get('start_time_formatted', '00:00')
                    confidence = segment.get('avg_confidence')
                    text = segment.get('text', '').strip()
                    
                    if confidence:
                        conf_display = f" *(confidence: {confidence*100:.1f}%)*"
                    else:
                        conf_display = ""
                    
                    md_lines.append(f"**[{timestamp}]** {text}{conf_display}")
                    md_lines.append("")
                
                if len(interim_segments) > 10:
                    md_lines.append(f"*... and {len(interim_segments) - 10} more interim segments*")
                    md_lines.append("")
            
        else:
            md_lines.append("## Transcript")
            md_lines.append("")
            md_lines.append("*No transcript available for this session.*")
            md_lines.append("")
        
        # Footer
        md_lines.append("---")
        md_lines.append("")
        md_lines.append("*Exported from Mina - Meeting Insights & Action Platform*")
        md_lines.append(f"*Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC*")
        md_lines.append(f"*Engine: OpenAI Whisper API*")
        
        return "\n".join(md_lines)
    
    @staticmethod
    def session_to_docx(session_id: int) -> Optional[io.BytesIO]:
        """
        Convert a session to DOCX format.
        
        Args:
            session_id: Database session ID
            
        Returns:
            BytesIO buffer containing DOCX data or None if session not found
        """
        # Get session detail
        session_detail = SessionService.get_session_detail(session_id)
        if not session_detail:
            return None
        
        session = session_detail['session']
        segments = session_detail['segments']
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading(session['title'], 0)
        title.alignment = 1  # Center alignment
        
        # Session info table
        doc.add_heading('Session Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Session ID', session['external_id']),
            ('Status', session['status'].title()),
            ('Started', session['started_at'][:16].replace('T', ' ') if session.get('started_at') else 'N/A'),
            ('Language', session.get('locale', 'Not specified'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
        
        # Transcript section
        doc.add_heading('Transcript', level=1)
        
        if segments:
            for segment in segments:
                # Add timestamp
                timestamp = segment.get('start_time', '')
                if timestamp:
                    timestamp_str = str(timestamp)[:8] if isinstance(timestamp, str) else 'N/A'
                else:
                    timestamp_str = 'N/A'
                
                # Add segment with formatting
                p = doc.add_paragraph()
                p.add_run(f"[{timestamp_str}] ").bold = True
                p.add_run(segment.get('text', ''))
        else:
            doc.add_paragraph('No transcript available for this session.')
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    @staticmethod
    def session_to_pdf(session_id: int) -> Optional[io.BytesIO]:
        """
        Convert a session to PDF format.
        
        Args:
            session_id: Database session ID
            
        Returns:
            BytesIO buffer containing PDF data or None if session not found
        """
        # Get session detail
        session_detail = SessionService.get_session_detail(session_id)
        if not session_detail:
            return None
        
        session = session_detail['session']
        segments = session_detail['segments']
        
        # Create PDF document
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            alignment=1,  # Center
            spaceAfter=30
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.darkblue,
            spaceAfter=12
        )
        
        # Build story
        story = []
        
        # Title
        story.append(Paragraph(session['title'], title_style))
        story.append(Spacer(1, 20))
        
        # Session information table
        story.append(Paragraph('Session Information', heading_style))
        
        info_data = [
            ['Session ID', session['external_id']],
            ['Status', session['status'].title()],
            ['Started', session['started_at'][:16].replace('T', ' ') if session.get('started_at') else 'N/A'],
            ['Language', session.get('locale', 'Not specified')],
            ['Segments', str(len(segments))]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Transcript section
        story.append(Paragraph('Transcript', heading_style))
        
        if segments:
            for segment in segments:
                # Format timestamp
                timestamp = segment.get('start_time', '')
                timestamp_str = str(timestamp)[:8] if isinstance(timestamp, str) else 'N/A'
                
                # Create segment text
                text = f"<b>[{timestamp_str}]</b> {segment.get('text', '')}"
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 8))
        else:
            story.append(Paragraph('No transcript available for this session.', styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer
    
    @staticmethod
    def get_export_filename(session_id: int, format_type: str = "md") -> str:
        """
        Generate appropriate filename for export.
        
        Args:
            session_id: Database session ID
            format_type: Export format ('md', 'txt', 'json', etc.)
            
        Returns:
            Suggested filename
        """
        # Get session for title
        session = SessionService.get_session_by_id(session_id)
        if session:
            # Sanitize title for filename
            safe_title = "".join(c for c in session.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '-').lower()[:30]  # Limit length
            return f"mina-{safe_title}-{session_id}.{format_type}"
        else:
            return f"mina-session-{session_id}.{format_type}"
    
    @staticmethod
    def session_to_txt(session_id: int) -> Optional[str]:
        """Convert a session to plain text format."""
        d = SessionService.get_session_detail(session_id)
        if not d: 
            return None
        finals = [s['text'].strip() for s in d['segments'] if s.get('is_final')]
        return "\n".join(finals)
    
    @staticmethod
    def session_to_vtt(session_id: int) -> Optional[str]:
        """Convert a session to WebVTT subtitle format."""
        d = SessionService.get_session_detail(session_id)
        if not d: 
            return None
        lines = ["WEBVTT", ""]
        for s in d['segments']:
            if not s.get('is_final'): 
                continue
            start = s.get('start_time_formatted','00:00.000').replace('.',',')
            end = s.get('end_time_formatted','00:00.000').replace('.',',')
            lines += [f"{start} --> {end}", s.get('text','').strip(), ""]
        return "\n".join(lines)


# Advanced Export System
class ExportFormat(Enum):
    """Supported export formats."""
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "markdown"
    HTML = "html"
    TXT = "txt"


class ExportTemplate(Enum):
    """Export template styles."""
    STANDARD = "standard"
    EXECUTIVE = "executive"
    TECHNICAL = "technical"
    MINIMAL = "minimal"
    BRANDED = "branded"


@dataclass
class ExportRequest:
    """Request data for export operations."""
    format: ExportFormat
    template: ExportTemplate
    session_ids: List[int]
    include_transcript: bool = True
    include_summary: bool = True
    include_tasks: bool = True
    include_analytics: bool = False
    custom_title: Optional[str] = None
    custom_header: Optional[str] = None
    custom_footer: Optional[str] = None
    date_range: Optional[Dict[str, str]] = None


@dataclass
class ExportResult:
    """Result of export operation."""
    success: bool
    file_path: Optional[str] = None
    file_content: Optional[bytes] = None
    filename: Optional[str] = None
    error_message: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


class AdvancedExportService:
    """Advanced export service that supports multiple sessions and templates."""
    
    def __init__(self):
        self.primary_color = HexColor('#6366f1')  # Indigo
        self.secondary_color = HexColor('#8b5cf6')  # Purple
        self.text_color = HexColor('#1f2937')  # Gray-800
        self.accent_color = HexColor('#10b981')  # Emerald
    
    def export_multiple_sessions(self, request: ExportRequest) -> ExportResult:
        """Export multiple sessions according to request."""
        try:
            if request.format == ExportFormat.PDF:
                return self._export_pdf_advanced(request)
            elif request.format == ExportFormat.DOCX:
                return self._export_docx_advanced(request)
            elif request.format == ExportFormat.MARKDOWN:
                return self._export_markdown_advanced(request)
            else:
                return ExportResult(
                    success=False,
                    error_message=f"Unsupported export format: {request.format}"
                )
                
        except Exception as e:
            logger.error(f"Advanced export failed: {e}")
            return ExportResult(
                success=False,
                error_message=f"Export failed: {str(e)}"
            )
    
    def get_session_data_advanced(self, session_ids: List[int]) -> List[Dict[str, Any]]:
        """Get comprehensive session data for export."""
        try:
            from models.session import Session
            from models.segment import Segment  
            from models.summary import Summary
            from models.task import Task
            from app import db
            
            sessions_data = []
            
            for session_id in session_ids:
                # Use existing SessionService for basic data
                session_detail = SessionService.get_session_detail(session_id)
                if not session_detail:
                    continue
                
                session = session_detail['session']
                segments = session_detail['segments']
                
                # Get additional data
                try:
                    # Get summary data
                    summary_query = db.session.query(Summary).filter_by(session_id=session_id).first()
                    summary_data = None
                    if summary_query:
                        summary_data = {
                            'content': summary_query.content,
                            'key_points': json.loads(summary_query.key_points) if summary_query.key_points else [],
                            'action_items': json.loads(summary_query.action_items) if summary_query.action_items else [],
                            'decisions': json.loads(summary_query.decisions) if summary_query.decisions else []
                        }
                    
                    # Get tasks
                    tasks_query = db.session.query(Task).filter_by(session_id=session_id).all()
                    tasks_data = [
                        {
                            'text': task.text,
                            'priority': task.priority,
                            'due_date': task.due_date,
                            'status': task.status,
                            'created_at': task.created_at
                        }
                        for task in tasks_query
                    ]
                    
                except Exception as e:
                    logger.warning(f"Could not get additional data for session {session_id}: {e}")
                    summary_data = None
                    tasks_data = []
                
                session_data = {
                    'id': session['id'],
                    'title': session['title'],
                    'external_id': session['external_id'],
                    'status': session['status'],
                    'started_at': session.get('started_at'),
                    'completed_at': session.get('completed_at'),
                    'locale': session.get('locale'),
                    'segments_count': session['segments_count'],
                    'transcript': [
                        {
                            'text': segment.get('text', ''),
                            'speaker': segment.get('speaker_label'),
                            'start_time': segment.get('start_time'),
                            'end_time': segment.get('end_time'),
                            'confidence': segment.get('avg_confidence'),
                            'is_final': segment.get('is_final', False)
                        }
                        for segment in segments
                    ],
                    'summary': summary_data,
                    'tasks': tasks_data,
                    'device_info': session.get('device_info', {})
                }
                
                sessions_data.append(session_data)
            
            return sessions_data
            
        except Exception as e:
            logger.error(f"Error getting advanced session data: {e}")
            return []
    
    def _export_pdf_advanced(self, request: ExportRequest) -> ExportResult:
        """Export to advanced PDF format."""
        try:
            sessions_data = self.get_session_data_advanced(request.session_ids)
            
            if not sessions_data:
                return ExportResult(
                    success=False,
                    error_message="No session data found for export"
                )
            
            # Create PDF document
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Setup styles
            styles = getSampleStyleSheet()
            self._setup_pdf_styles(styles, request.template)
            
            # Build document content
            story = []
            
            # Title page
            self._add_pdf_title_page(story, request, sessions_data, styles)
            
            # Table of contents (for multi-session exports)
            if len(sessions_data) > 1:
                self._add_pdf_table_of_contents(story, sessions_data, styles)
            
            # Session content
            for i, session_data in enumerate(sessions_data):
                if i > 0:
                    story.append(PageBreak())
                self._add_pdf_session_content(story, session_data, request, styles)
            
            # Build PDF
            doc.build(story)
            
            # Generate filename
            if len(sessions_data) == 1:
                filename = f"mina_meeting_{sessions_data[0]['id']}_{datetime.now().strftime('%Y%m%d')}.pdf"
            else:
                filename = f"mina_meetings_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            return ExportResult(
                success=True,
                file_content=buffer.getvalue(),
                filename=filename,
                metadata={
                    'format': 'pdf',
                    'template': request.template.value,
                    'sessions_count': len(sessions_data)
                }
            )
            
        except Exception as e:
            logger.error(f"Error exporting to advanced PDF: {e}")
            return ExportResult(
                success=False,
                error_message=f"PDF export failed: {str(e)}"
            )
    
    def _setup_pdf_styles(self, styles, template: ExportTemplate):
        """Setup PDF styles based on template."""
        # Title style
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=self.primary_color,
            spaceAfter=20,
            alignment=TA_CENTER
        ))
        
        # Section header style
        styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=self.primary_color,
            spaceAfter=10,
            spaceBefore=15
        ))
        
        # Speaker style
        styles.add(ParagraphStyle(
            name='Speaker',
            parent=styles['Normal'],
            fontSize=11,
            textColor=self.secondary_color,
            fontName='Helvetica-Bold',
            spaceAfter=5
        ))
        
        # Template-specific adjustments
        if template == ExportTemplate.EXECUTIVE:
            styles['CustomTitle'].fontSize = 28
            styles['CustomTitle'].textColor = black
        elif template == ExportTemplate.TECHNICAL:
            styles.add(ParagraphStyle(
                name='Code',
                parent=styles['Normal'],
                fontName='Courier',
                fontSize=10,
                textColor=self.text_color
            ))
    
    def _add_pdf_title_page(self, story: List, request: ExportRequest, sessions_data: List[Dict], styles):
        """Add title page to PDF."""
        # Main title
        title = request.custom_title or f"Meeting Export - {datetime.now().strftime('%B %d, %Y')}"
        story.append(Paragraph(title, styles['CustomTitle']))
        story.append(Spacer(1, 30))
        
        # Subtitle
        if len(sessions_data) == 1:
            subtitle = f"Meeting: {sessions_data[0]['title']}"
        else:
            subtitle = f"{len(sessions_data)} Meetings"
        
        story.append(Paragraph(subtitle, styles['Heading2']))
        story.append(Spacer(1, 20))
        
        # Generation info
        story.append(Paragraph("Generated by Mina Meeting Intelligence", styles['Normal']))
        story.append(Paragraph(f"Export Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Summary table
        summary_data = [
            ['Sessions Included', str(len(sessions_data))],
            ['Export Format', request.format.value.upper()],
            ['Template', request.template.value.title()],
            ['Includes Transcript', 'Yes' if request.include_transcript else 'No'],
            ['Includes Summary', 'Yes' if request.include_summary else 'No'],
            ['Includes Tasks', 'Yes' if request.include_tasks else 'No']
        ]
        
        summary_table = Table(summary_data, colWidths=[2*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f8fafc')),
            ('TEXTCOLOR', (0, 0), (-1, -1), self.text_color),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, HexColor('#e5e7eb'))
        ]))
        
        story.append(summary_table)
        story.append(PageBreak())
    
    def _add_pdf_table_of_contents(self, story: List, sessions_data: List[Dict], styles):
        """Add table of contents for multi-session exports."""
        story.append(Paragraph("Table of Contents", styles['SectionHeader']))
        story.append(Spacer(1, 20))
        
        for i, session in enumerate(sessions_data):
            date_str = session.get('started_at', '')[:10] if session.get('started_at') else 'No Date'
            story.append(Paragraph(
                f"{i+1}. {session['title']} - {date_str}",
                styles['Normal']
            ))
        
        story.append(PageBreak())
    
    def _add_pdf_session_content(self, story: List, session_data: Dict, request: ExportRequest, styles):
        """Add session content to PDF."""
        # Session header
        story.append(Paragraph(session_data['title'], styles['SectionHeader']))
        
        if session_data.get('started_at'):
            date_str = session_data['started_at'][:16].replace('T', ' ')
            story.append(Paragraph(f"Date: {date_str}", styles['Normal']))
        
        story.append(Paragraph(f"Session ID: {session_data['external_id']}", styles['Normal']))
        story.append(Paragraph(f"Status: {session_data['status'].title()}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Summary section
        if request.include_summary and session_data['summary']:
            story.append(Paragraph("Executive Summary", styles['SectionHeader']))
            
            if session_data['summary']['content']:
                story.append(Paragraph(session_data['summary']['content'], styles['Normal']))
                story.append(Spacer(1, 15))
            
            # Key points
            if session_data['summary']['key_points']:
                story.append(Paragraph("Key Points:", styles['Heading3']))
                for point in session_data['summary']['key_points']:
                    story.append(Paragraph(f"• {point}", styles['Normal']))
                story.append(Spacer(1, 10))
            
            # Action items
            if session_data['summary']['action_items']:
                story.append(Paragraph("Action Items:", styles['Heading3']))
                for item in session_data['summary']['action_items']:
                    story.append(Paragraph(f"• {item}", styles['Normal']))
                story.append(Spacer(1, 10))
            
            # Decisions
            if session_data['summary']['decisions']:
                story.append(Paragraph("Key Decisions:", styles['Heading3']))
                for decision in session_data['summary']['decisions']:
                    story.append(Paragraph(f"• {decision}", styles['Normal']))
                story.append(Spacer(1, 15))
        
        # Tasks section
        if request.include_tasks and session_data['tasks']:
            story.append(Paragraph("Tasks", styles['SectionHeader']))
            
            for task in session_data['tasks']:
                task_text = f"• {task['text']}"
                if task['due_date']:
                    task_text += f" (Due: {task['due_date'].strftime('%m/%d/%Y')})"
                if task['priority']:
                    task_text += f" [{task['priority'].upper()}]"
                
                story.append(Paragraph(task_text, styles['Normal']))
            
            story.append(Spacer(1, 15))
        
        # Transcript section
        if request.include_transcript and session_data['transcript']:
            story.append(Paragraph("Full Transcript", styles['SectionHeader']))
            
            final_segments = [s for s in session_data['transcript'] if s.get('is_final')]
            if not final_segments:
                final_segments = session_data['transcript'][:10]  # Limit if no final segments
            
            current_speaker = None
            for segment in final_segments:
                if segment['speaker'] != current_speaker:
                    current_speaker = segment['speaker']
                    speaker_name = segment['speaker'] or "Speaker"
                    story.append(Paragraph(f"{speaker_name}:", styles['Speaker']))
                
                story.append(Paragraph(segment['text'], styles['Normal']))
                story.append(Spacer(1, 5))
    
    def _export_markdown_advanced(self, request: ExportRequest) -> ExportResult:
        """Export to advanced Markdown format."""
        try:
            sessions_data = self.get_session_data_advanced(request.session_ids)
            
            if not sessions_data:
                return ExportResult(
                    success=False,
                    error_message="No session data found for export"
                )
            
            # Build markdown content
            markdown_content = []
            
            # Document header
            title = request.custom_title or f"Meeting Export - {datetime.now().strftime('%B %d, %Y')}"
            markdown_content.append(f"# {title}\n")
            
            if len(sessions_data) == 1:
                subtitle = f"Meeting: {sessions_data[0]['title']}"
            else:
                subtitle = f"{len(sessions_data)} Meetings"
            
            markdown_content.append(f"## {subtitle}\n")
            markdown_content.append("Generated by Mina Meeting Intelligence  ")
            markdown_content.append(f"Export Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
            
            # Export info
            markdown_content.append("### Export Information\n")
            markdown_content.append(f"- **Sessions Included:** {len(sessions_data)}")
            markdown_content.append(f"- **Export Format:** {request.format.value.upper()}")
            markdown_content.append(f"- **Template:** {request.template.value.title()}")
            markdown_content.append(f"- **Includes Transcript:** {'Yes' if request.include_transcript else 'No'}")
            markdown_content.append(f"- **Includes Summary:** {'Yes' if request.include_summary else 'No'}")
            markdown_content.append(f"- **Includes Tasks:** {'Yes' if request.include_tasks else 'No'}\n")
            
            # Table of contents for multi-session exports
            if len(sessions_data) > 1:
                markdown_content.append("## Table of Contents\n")
                for i, session in enumerate(sessions_data):
                    markdown_content.append(f"{i+1}. [{session['title']}](#{session['title'].lower().replace(' ', '-')})")
                markdown_content.append("")
            
            # Session content
            for session_data in sessions_data:
                self._add_markdown_session_content(markdown_content, session_data, request)
            
            # Join content
            final_content = "\n".join(markdown_content)
            
            # Generate filename
            if len(sessions_data) == 1:
                filename = f"mina_meeting_{sessions_data[0]['id']}_{datetime.now().strftime('%Y%m%d')}.md"
            else:
                filename = f"mina_meetings_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            
            return ExportResult(
                success=True,
                file_content=final_content.encode('utf-8'),
                filename=filename,
                metadata={
                    'format': 'markdown',
                    'template': request.template.value,
                    'sessions_count': len(sessions_data),
                    'word_count': len(final_content.split())
                }
            )
            
        except Exception as e:
            logger.error(f"Error exporting to advanced Markdown: {e}")
            return ExportResult(
                success=False,
                error_message=f"Markdown export failed: {str(e)}"
            )
    
    def _add_markdown_session_content(self, content: List[str], session_data: Dict, request: ExportRequest):
        """Add session content to Markdown."""
        # Session header
        content.append(f"## {session_data['title']}\n")
        
        if session_data.get('started_at'):
            date_str = session_data['started_at'][:16].replace('T', ' ')
            content.append(f"**Date:** {date_str}")
        
        content.append(f"**Session ID:** {session_data['external_id']}")
        content.append(f"**Status:** {session_data['status'].title()}")
        content.append("")
        
        # Summary section
        if request.include_summary and session_data['summary']:
            content.append("### Executive Summary\n")
            
            if session_data['summary']['content']:
                content.append(session_data['summary']['content'])
                content.append("")
            
            # Key points
            if session_data['summary']['key_points']:
                content.append("#### Key Points\n")
                for point in session_data['summary']['key_points']:
                    content.append(f"- {point}")
                content.append("")
            
            # Action items
            if session_data['summary']['action_items']:
                content.append("#### Action Items\n")
                for item in session_data['summary']['action_items']:
                    content.append(f"- {item}")
                content.append("")
            
            # Decisions
            if session_data['summary']['decisions']:
                content.append("#### Key Decisions\n")
                for decision in session_data['summary']['decisions']:
                    content.append(f"- {decision}")
                content.append("")
        
        # Tasks section
        if request.include_tasks and session_data['tasks']:
            content.append("### Tasks\n")
            
            for task in session_data['tasks']:
                task_text = f"- {task['text']}"
                if task['due_date']:
                    task_text += f" *(Due: {task['due_date'].strftime('%m/%d/%Y')})*"
                if task['priority']:
                    task_text += f" **[{task['priority'].upper()}]**"
                
                content.append(task_text)
            
            content.append("")
        
        # Transcript section
        if request.include_transcript and session_data['transcript']:
            content.append("### Full Transcript\n")
            
            final_segments = [s for s in session_data['transcript'] if s.get('is_final')]
            if not final_segments:
                final_segments = session_data['transcript'][:20]  # Limit if no final segments
            
            current_speaker = None
            for segment in final_segments:
                if segment['speaker'] != current_speaker:
                    current_speaker = segment['speaker']
                    speaker_name = segment['speaker'] or "Speaker"
                    content.append(f"**{speaker_name}:**")
                
                content.append(segment['text'])
                content.append("")
    
    def _export_docx_advanced(self, request: ExportRequest) -> ExportResult:
        """Export to advanced DOCX format."""
        try:
            sessions_data = self.get_session_data_advanced(request.session_ids)
            
            if not sessions_data:
                return ExportResult(
                    success=False,
                    error_message="No session data found for export"
                )
            
            # Create Word document
            doc = Document()
            
            # Add content
            self._add_docx_header(doc, request, sessions_data)
            
            for i, session_data in enumerate(sessions_data):
                if i > 0:
                    doc.add_page_break()
                self._add_docx_session_content(doc, session_data, request)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            
            # Generate filename
            if len(sessions_data) == 1:
                filename = f"mina_meeting_{sessions_data[0]['id']}_{datetime.now().strftime('%Y%m%d')}.docx"
            else:
                filename = f"mina_meetings_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            return ExportResult(
                success=True,
                file_content=buffer.getvalue(),
                filename=filename,
                metadata={
                    'format': 'docx',
                    'template': request.template.value,
                    'sessions_count': len(sessions_data)
                }
            )
            
        except Exception as e:
            logger.error(f"Error exporting to advanced DOCX: {e}")
            return ExportResult(
                success=False,
                error_message=f"DOCX export failed: {str(e)}"
            )
    
    def _add_docx_header(self, doc: Document, request: ExportRequest, sessions_data: List[Dict]):
        """Add document header."""
        # Title
        title = request.custom_title or f"Meeting Export - {datetime.now().strftime('%B %d, %Y')}"
        title_paragraph = doc.add_paragraph(title)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        if len(sessions_data) == 1:
            subtitle = f"Meeting: {sessions_data[0]['title']}"
        else:
            subtitle = f"{len(sessions_data)} Meetings"
        
        doc.add_paragraph(subtitle, style='Heading 2')
        
        # Generation info
        doc.add_paragraph("Generated by Mina Meeting Intelligence")
        doc.add_paragraph(f"Export Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph()  # Empty line
    
    def _add_docx_session_content(self, doc: Document, session_data: Dict, request: ExportRequest):
        """Add session content to DOCX."""
        # Session header
        doc.add_paragraph(session_data['title'], style='Heading 1')
        
        if session_data.get('started_at'):
            date_str = session_data['started_at'][:16].replace('T', ' ')
            doc.add_paragraph(f"Date: {date_str}")
        
        doc.add_paragraph(f"Session ID: {session_data['external_id']}")
        doc.add_paragraph()  # Empty line
        
        # Summary section
        if request.include_summary and session_data['summary']:
            doc.add_paragraph("Executive Summary", style='Heading 2')
            
            if session_data['summary']['content']:
                doc.add_paragraph(session_data['summary']['content'])
            
            # Key points
            if session_data['summary']['key_points']:
                doc.add_paragraph("Key Points:", style='Heading 3')
                for point in session_data['summary']['key_points']:
                    doc.add_paragraph(f"• {point}")
            
            # Action items
            if session_data['summary']['action_items']:
                doc.add_paragraph("Action Items:", style='Heading 3')
                for item in session_data['summary']['action_items']:
                    doc.add_paragraph(f"• {item}")
            
            # Decisions
            if session_data['summary']['decisions']:
                doc.add_paragraph("Key Decisions:", style='Heading 3')
                for decision in session_data['summary']['decisions']:
                    doc.add_paragraph(f"• {decision}")
        
        # Tasks section
        if request.include_tasks and session_data['tasks']:
            doc.add_paragraph("Tasks", style='Heading 2')
            
            for task in session_data['tasks']:
                task_text = f"• {task['text']}"
                if task['due_date']:
                    task_text += f" (Due: {task['due_date'].strftime('%m/%d/%Y')})"
                if task['priority']:
                    task_text += f" [{task['priority'].upper()}]"
                
                doc.add_paragraph(task_text)
        
        # Transcript section
        if request.include_transcript and session_data['transcript']:
            doc.add_paragraph("Full Transcript", style='Heading 2')
            
            final_segments = [s for s in session_data['transcript'] if s.get('is_final')]
            if not final_segments:
                final_segments = session_data['transcript'][:20]  # Limit if no final segments
            
            current_speaker = None
            for segment in final_segments:
                if segment['speaker'] != current_speaker:
                    current_speaker = segment['speaker']
                    speaker_name = segment['speaker'] or "Speaker"
                    p = doc.add_paragraph()
                    p.add_run(f"{speaker_name}:").bold = True
                
                doc.add_paragraph(segment['text'])


# Create global instance of advanced export service
advanced_export_service = AdvancedExportService()